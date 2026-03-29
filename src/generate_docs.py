from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
import os

# Get absolute path to project root
ROOT_DIR = os.path.abspath(os.path.dirname(os.path.abspath(__file__)) + "/..")
PDF_DIR = os.path.join(ROOT_DIR, "data", "raw", "pdf")
WORD_DIR = os.path.join(ROOT_DIR, "data", "raw", "word")
EXCEL_DIR = os.path.join(ROOT_DIR, "data", "raw", "excel")

print(f"Root dir: {ROOT_DIR}")
print(f"PDF dir: {PDF_DIR}")

def generate_normal_pdf():
    os.makedirs(PDF_DIR, exist_ok=True)
    filepath = os.path.join(PDF_DIR, "automotive_report_normal.pdf")
    doc = SimpleDocTemplate(filepath, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("Automotive Industry Recall Report 2020", styles["Title"]))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph("Introduction", styles["Heading1"]))
    elements.append(Paragraph(
        "This report provides an overview of vehicle recalls issued in 2020 for major "
        "automotive brands including Audi, Volkswagen, and Skoda. Data sourced from NHTSA.",
        styles["Normal"]
    ))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("Recall Summary Table (with borders)", styles["Heading2"]))
    data = [
        ["Make", "Model", "Year", "Recalls"],
        ["Audi", "A4", "2020", "3"],
        ["Volkswagen", "Golf", "2020", "2"],
        ["Volkswagen", "Passat", "2020", "3"],
    ]
    table = Table(data)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 1, colors.black),
        ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("Recall Summary Table (without borders)", styles["Heading2"]))
    table2 = Table(data)
    table2.setStyle(TableStyle([
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightblue),
    ]))
    elements.append(table2)

    doc.build(elements)
    print(f"Generated: {filepath}")

def generate_twocolumn_pdf():
    os.makedirs(PDF_DIR, exist_ok=True)
    filepath = os.path.join(PDF_DIR, "automotive_report_twocolumn.pdf")
    doc = SimpleDocTemplate(filepath, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("Automotive Two-Column Report", styles["Title"]))
    elements.append(Spacer(1, 12))

    col_data = [
        [
            Paragraph("<b>Audi A4 Recalls</b><br/>The Audi A4 2020 had 3 recalls related "
                     "to airbag, brake, and engine issues reported to NHTSA.", styles["Normal"]),
            Paragraph("<b>Volkswagen Golf Recalls</b><br/>The Volkswagen Golf 2020 had 2 recalls "
                     "related to electrical systems and fuel systems.", styles["Normal"]),
        ]
    ]
    two_col_table = Table(col_data, colWidths=[250, 250])
    two_col_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
    ]))
    elements.append(two_col_table)
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("Component Recall Details (with borders)", styles["Heading2"]))
    data = [
        ["Component", "Issue", "Affected Units"],
        ["Airbag", "Inflator rupture", "15,000"],
        ["Brakes", "Brake fluid leak", "8,200"],
        ["Engine", "Oil pressure loss", "5,100"],
    ]
    table = Table(data)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.darkblue),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 1, colors.black),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("Component Recall Details (without borders)", styles["Heading2"]))
    table2 = Table(data)
    table2.setStyle(TableStyle([
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.darkblue),
    ]))
    elements.append(table2)

    doc.build(elements)
    print(f"Generated: {filepath}")

def generate_word_doc():
    os.makedirs(WORD_DIR, exist_ok=True)
    filepath = os.path.join(WORD_DIR, "automotive_report.docx")
    doc = Document()

    doc.add_heading("Automotive Industry Recall Report 2020", 0)
    doc.add_paragraph(
        "This report provides an overview of vehicle recalls issued in 2020. "
        "Data sourced from NHTSA database covering major VW Group vehicles."
    )

    doc.add_heading("Recall Overview", level=1)
    doc.add_paragraph(
        "In 2020, several VW Group vehicles were subject to safety recalls. "
        "The recalls covered issues ranging from airbag defects to engine problems."
    )

    doc.add_heading("Table with Borders", level=2)
    table1 = doc.add_table(rows=4, cols=4)
    table1.style = "Table Grid"
    headers = ["Make", "Model", "Year", "Recalls"]
    rows_data = [
        ["Audi", "A4", "2020", "3"],
        ["Volkswagen", "Golf", "2020", "2"],
        ["Volkswagen", "Passat", "2020", "3"],
    ]
    for i, header in enumerate(headers):
        table1.rows[0].cells[i].text = header
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, value in enumerate(row_data):
            table1.rows[row_idx + 1].cells[col_idx].text = value

    doc.add_paragraph("")
    doc.add_heading("Table without Borders", level=2)
    table2 = doc.add_table(rows=4, cols=4)
    for i, header in enumerate(headers):
        table2.rows[0].cells[i].text = header
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, value in enumerate(row_data):
            table2.rows[row_idx + 1].cells[col_idx].text = value

    doc.add_heading("Two-Column Section", level=1)
    doc.add_paragraph("Left Column: Audi A4 had 3 recalls in 2020 including airbag and brake issues.")
    doc.add_paragraph("Right Column: Volkswagen Golf had 2 recalls related to electrical systems.")

    doc.save(filepath)
    print(f"Generated: {filepath}")

def generate_excel():
    os.makedirs(EXCEL_DIR, exist_ok=True)
    filepath = os.path.join(EXCEL_DIR, "automotive_data.xlsx")
    wb = openpyxl.Workbook()

    ws1 = wb.active
    ws1.title = "Recall Summary"
    headers = ["Make", "Model", "Year", "Recall Count", "Affected Units", "Total Cost (USD)"]
    ws1.append(headers)

    data = [
        ["Audi", "A4", 2020, 3, 15000, 2500000],
        ["Volkswagen", "Golf", 2020, 2, 8200, 1200000],
        ["Volkswagen", "Passat", 2020, 3, 5100, 980000],
    ]
    for row in data:
        ws1.append(row)

    ws1.append(["TOTAL", "", "", "=SUM(D2:D4)", "=SUM(E2:E4)", "=SUM(F2:F4)"])

    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    for cell in ws1[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    thin = Side(style="thin")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for row in ws1.iter_rows(min_row=1, max_row=5, min_col=1, max_col=6):
        for cell in row:
            cell.border = border

    ws2 = wb.create_sheet("Component Details")
    ws2.append(["Component", "Issue", "Make", "Model", "Affected Units"])
    ws2.append(["Airbag", "Inflator rupture", "Audi", "A4", 15000])
    ws2.append(["Brakes", "Brake fluid leak", "Volkswagen", "Golf", 8200])
    ws2.append(["Engine", "Oil pressure loss", "Volkswagen", "Passat", 5100])

    wb.save(filepath)
    print(f"Generated: {filepath}")

if __name__ == "__main__":
    generate_normal_pdf()
    generate_twocolumn_pdf()
    generate_word_doc()
    generate_excel()
    print("All documents generated successfully!")