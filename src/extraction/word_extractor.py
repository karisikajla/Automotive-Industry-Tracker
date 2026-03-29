import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..')))

from docx import Document
import chardet
from datetime import datetime
from src.utils.logger import logging
from src.storage.mongo import save_to_mongo

def detect_encoding(filepath):
    with open(filepath, "rb") as f:
        raw = f.read()
    result = chardet.detect(raw)
    encoding = result.get("encoding", "utf-8")
    logging.info(f"Detected encoding for {filepath}: {encoding}")
    return encoding

def extract_word(filepath):
    logging.info(f"Starting Word extraction: {filepath}")
    result = {
        "file_name": os.path.basename(filepath),
        "document_type": "word",
        "source": filepath,
        "extraction_timestamp": datetime.utcnow().isoformat(),
        "extraction_library": "python-docx",
        "paragraphs": [],
        "tables": []
    }

    try:
        detect_encoding(filepath)
        doc = Document(filepath)

        for para in doc.paragraphs:
            if para.text.strip():
                result["paragraphs"].append({
                    "style": para.style.name,
                    "text": para.text.strip()
                })

        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            result["tables"].append({
                "table_index": table_idx,
                "data": table_data
            })
            logging.info(f"Extracted table {table_idx} from {os.path.basename(filepath)}")

        logging.info(f"Extracted {len(result['paragraphs'])} paragraphs from {os.path.basename(filepath)}")
        save_to_mongo(result, "word_extraction")
        logging.info(f"Word extraction complete: {filepath}")

    except Exception as e:
        logging.error(f"Error extracting Word {filepath}: {e}")

    return result

def extract_all_word_docs():
    ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..'))
    word_dir = os.path.join(ROOT_DIR, "data", "raw", "word")
    results = []
    for filename in os.listdir(word_dir):
        if filename.endswith(".docx"):
            filepath = os.path.join(word_dir, filename)
            result = extract_word(filepath)
            results.append(result)
    return results

results = extract_all_word_docs()
for r in results:
    print(f"Extracted {r['file_name']}: {len(r['paragraphs'])} paragraphs, {len(r['tables'])} tables")